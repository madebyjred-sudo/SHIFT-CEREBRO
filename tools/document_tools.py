"""
Document Generation Tools for Shift AI Gateway
Based on python-docx: https://github.com/python-openxml/python-docx
Installation: pip install python-docx

This module provides LangChain-compatible tools for generating
Word documents (.docx) from agent responses.
"""

import os
import tempfile
from datetime import datetime
from typing import Optional, List, Dict, Any
from langchain_core.tools import tool
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Local storage configuration
DOCUMENTS_DIR = os.path.join(os.path.dirname(__file__), "..", "generated_documents")
os.makedirs(DOCUMENTS_DIR, exist_ok=True)


def _set_cell_shading(cell, color: str):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_heading_with_style(doc: Document, text: str, level: int = 1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri'
        run.font.bold = True
        if level == 0:
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(46, 116, 181)  # Medium blue
        else:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(68, 114, 196)  # Light blue
    return heading


def _add_formatted_paragraph(doc: Document, text: str, bold: bool = False, 
                            italic: bool = False, font_size: int = 11):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return para


@tool
def create_word_document(
    title: str,
    content: str,
    subtitle: Optional[str] = None,
    author: str = "Shift AI",
    sections: Optional[List[Dict[str, str]]] = None,
    filename: Optional[str] = None
) -> str:
    """
    Create a professional Word document (.docx) with formatted content.
    
    This tool generates a properly formatted Microsoft Word document
    using the python-docx library. The document includes professional
    styling, headers, and structured content.
    
    Args:
        title: Main title of the document (required)
        content: Main content text (supports basic markdown-like formatting)
        subtitle: Optional subtitle displayed under the title
        author: Document author name (default: "Shift AI")
        sections: Optional list of sections, each with 'heading' and 'content' keys
                 Example: [{"heading": "Section 1", "content": "Content here"}]
        filename: Optional custom filename (without extension)
                 If not provided, uses sanitized title
    
    Returns:
        str: URL/path to the generated document file
    
    Example:
        create_word_document(
            title="Marketing Brief Q1 2026",
            subtitle="Campaign Strategy Overview",
            content="This document outlines our marketing strategy...",
            author="Carmen - CEO",
            sections=[
                {"heading": "Objectives", "content": "Increase brand awareness by 25%"},
                {"heading": "Target Audience", "content": "B2B SaaS decision makers"}
            ]
        )
    """
    try:
        # Create document
        doc = Document()
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 121)  # Professional dark blue
        
        # Add subtitle if provided
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle_para.add_run(subtitle)
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            run.italic = True
        
        # Add spacing
        doc.add_paragraph()
        
        # Add metadata (author and date)
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current_date = datetime.now().strftime("%d de %B de %Y")
        meta_text = f"Autor: {author} | Fecha: {current_date}"
        run = meta_para.add_run(meta_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add horizontal line (using paragraph border)
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator_run = separator.add_run("─" * 50)
        separator_run.font.color.rgb = RGBColor(200, 200, 200)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Add main content
        if content:
            # Process content line by line for basic formatting
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    text = line.strip('*')
                    _add_formatted_paragraph(doc, text, bold=True)
                elif line.startswith('*') and line.endswith('*'):
                    # Italic text
                    text = line.strip('*')
                    _add_formatted_paragraph(doc, text, italic=True)
                elif line.startswith('- '):
                    # Bullet point
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(line[2:])
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                else:
                    # Regular paragraph
                    _add_formatted_paragraph(doc, line)
        
        # Add sections if provided
        if sections:
            for section in sections:
                # Add spacing before section
                doc.add_paragraph()
                
                # Add section heading
                heading_text = section.get('heading', '')
                if heading_text:
                    _add_heading_with_style(doc, heading_text, level=1)
                
                # Add section content
                section_content = section.get('content', '')
                if section_content:
                    lines = section_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            doc.add_paragraph()
                        elif line.startswith('- '):
                            para = doc.add_paragraph(style='List Bullet')
                            run = para.add_run(line[2:])
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                        else:
                            _add_formatted_paragraph(doc, line)
        
        # Add footer with branding
        doc.add_paragraph()
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("─  Generado por Shift AI  ─")
        footer_run.font.name = 'Calibri'
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(180, 180, 180)
        footer_run.italic = True
        
        # Generate filename
        if not filename:
            # Sanitize title for filename
            filename = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' 
                             for c in title).strip().replace(' ', '_')
        
        # Ensure unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        final_filename = f"{filename}_{timestamp}.docx"
        file_path = os.path.join(DOCUMENTS_DIR, final_filename)
        
        # Save document
        doc.save(file_path)
        
        # Return relative URL path (will be served via static endpoint)
        return f"/documents/{final_filename}"
        
    except Exception as e:
        return f"Error creating document: {str(e)}"


@tool
def create_brief_document(
    project_name: str,
    objectives: str,
    target_audience: str,
    key_messages: str,
    timeline: str,
    budget: Optional[str] = None,
    author: str = "Shift AI"
) -> str:
    """
    Create a structured marketing/creative brief document.
    
    Specialized tool for creating professional briefs with
    predefined sections for marketing projects.
    
    Args:
        project_name: Name of the project/campaign
        objectives: Project objectives and goals
        target_audience: Description of target audience
        key_messages: Key messages and positioning
        timeline: Project timeline and milestones
        budget: Optional budget information
        author: Document author
    
    Returns:
        str: URL/path to the generated brief document
    """
    sections = [
        {"heading": "1. Objetivos del Proyecto", "content": objectives},
        {"heading": "2. Audiencia Objetivo", "content": target_audience},
        {"heading": "3. Mensajes Clave", "content": key_messages},
        {"heading": "4. Timeline", "content": timeline},
    ]
    
    if budget:
        sections.append({"heading": "5. Presupuesto", "content": budget})
    
    return create_word_document(
        title=f"Brief: {project_name}",
        subtitle="Documento de Estrategia y Dirección",
        content="",
        author=author,
        sections=sections,
        filename=f"brief_{project_name.lower().replace(' ', '_')}"
    )


@tool
def create_meeting_minutes(
    meeting_title: str,
    attendees: str,
    agenda: str,
    discussion_points: str,
    action_items: str,
    next_steps: str,
    author: str = "Shift AI"
) -> str:
    """
    Create a professional meeting minutes document.
    
    Args:
        meeting_title: Title/subject of the meeting
        attendees: List of attendees
        agenda: Meeting agenda items
        discussion_points: Summary of discussions
        action_items: Action items with owners
        next_steps: Next steps and follow-ups
        author: Document author
    
    Returns:
        str: URL/path to the generated document
    """
    current_date = datetime.now().strftime("%d de %B de %Y")
    
    sections = [
        {"heading": "Información General", 
         "content": f"Fecha: {current_date}\nAsistentes: {attendees}\n"},
        {"heading": "Agenda", "content": agenda},
        {"heading": "Puntos de Discusión", "content": discussion_points},
        {"heading": "Action Items", "content": action_items},
        {"heading": "Próximos Pasos", "content": next_steps},
    ]
    
    return create_word_document(
        title=f"Acta de Reunión: {meeting_title}",
        content="",
        author=author,
        sections=sections,
        filename=f"acta_{meeting_title.lower().replace(' ', '_')}"
    )


# Export tools list for LangChain integration
DOCUMENT_TOOLS = [
    create_word_document,
    create_brief_document,
    create_meeting_minutes,
]